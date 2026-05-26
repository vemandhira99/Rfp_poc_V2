import json
import sys
from pathlib import Path

import requests
from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))

from app.core.database import SessionLocal
from app.models.rfp import RFPDraftSection
from app.services.local_embedding_service import check_embedding_model
from app.services.local_llm_service import check_ollama_health


API = "http://127.0.0.1:8001"
SCRATCH = Path("scratch")


def main() -> None:
    SCRATCH.mkdir(exist_ok=True)
    evidence = {}

    health = requests.get(f"{API}/health", timeout=20).json()
    evidence["health"] = health
    evidence["ollama"] = check_ollama_health()
    evidence["embedding_model"] = check_embedding_model()

    tiny = SCRATCH / "phase2_tiny.txt"
    tiny.write_text("Hello RFP.", encoding="utf-8")
    tiny_upload = upload_file(tiny)
    evidence["tiny_upload"] = tiny_upload

    valid_txt = SCRATCH / "phase2_valid_rfp.txt"
    valid_txt.write_text(valid_rfp_text(), encoding="utf-8")
    valid_upload = upload_file(valid_txt)
    evidence["valid_upload"] = valid_upload
    valid_id = valid_upload["rfp_id"]

    docx_path = SCRATCH / "phase2_searchable_rfp.docx"
    create_docx(docx_path)
    docx_upload = upload_file(docx_path)
    evidence["docx_upload"] = docx_upload

    embed = requests.post(f"{API}/rfps/{valid_id}/embed", timeout=240).json()
    rfp_after_embed = requests.get(f"{API}/rfps/{valid_id}", timeout=30).json()
    evidence["embedding"] = {"embed_response": embed, "rfp_after_embed": rfp_after_embed}

    retrieval = requests.get(
        f"{API}/rfps/{valid_id}/retrieval-test",
        params={"q": "scope requirements deadline"},
        timeout=240,
    ).json()
    evidence["retrieval"] = {
        "count": len(retrieval.get("results", [])),
        "first_type": retrieval.get("results", [{}])[0].get("retrieval_type") if retrieval.get("results") else None,
        "first_score": retrieval.get("results", [{}])[0].get("score") if retrieval.get("results") else None,
    }

    no_embed_upload = upload_file(SCRATCH / "phase2_valid_rfp.txt")
    lexical_fallback = requests.get(
        f"{API}/rfps/{no_embed_upload['rfp_id']}/retrieval-test",
        params={"q": "scope requirements deadline"},
        timeout=60,
    ).json()
    evidence["lexical_fallback"] = {
        "rfp_id": no_embed_upload["rfp_id"],
        "count": len(lexical_fallback.get("results", [])),
        "first_type": lexical_fallback.get("results", [{}])[0].get("retrieval_type") if lexical_fallback.get("results") else None,
    }

    evidence["chat_hi"] = chat(valid_id, "hi")
    evidence["chat_about"] = summarize_chat(chat(valid_id, "What is this RFP about?"))
    evidence["chat_requirements"] = summarize_chat(chat(valid_id, "What are the main requirements?"))
    evidence["chat_deadline"] = summarize_chat(chat(valid_id, "What is the deadline?"))

    draft = requests.post(f"{API}/private-rfp/{valid_id}/generate-short-draft", timeout=1500).json()
    sections = requests.get(f"{API}/private-rfp/{valid_id}/draft", timeout=60).json()
    evidence["draft_generation"] = {
        "response": draft,
        "section_count": len(sections),
        "orders": [section["section_order"] for section in sections],
        "duplicate_headings": duplicate_heading_count(sections),
    }

    quality = requests.get(f"{API}/private-rfp/{valid_id}/quality", timeout=60).json()
    evidence["quality"] = {
        "overall_status": quality.get("overall_status"),
        "check_names": [check.get("name") for check in quality.get("checks", [])],
        "check_count": len(quality.get("checks", [])),
    }

    export = requests.post(f"{API}/private-rfp/{valid_id}/export-docx", timeout=120).json()
    export_path = Path(export["file_path"])
    if not export_path.is_absolute():
        export_path = Path.cwd() / export_path
    download = requests.get(f"{API}/private-rfp/{valid_id}/download", timeout=120)
    evidence["export"] = {
        "response": export,
        "exists": export_path.exists(),
        "size": export_path.stat().st_size if export_path.exists() else 0,
        "download_status": download.status_code,
        "download_content_type": download.headers.get("content-type"),
        "privacy_note_found": docx_contains(export_path, "Generated locally. No external AI provider used.") if export_path.exists() else False,
    }

    db = SessionLocal()
    try:
        stored_sections = db.query(RFPDraftSection).filter(RFPDraftSection.rfp_id == valid_id).order_by(RFPDraftSection.section_order).all()
        evidence["db_sections"] = {"count": len(stored_sections), "orders": [section.section_order for section in stored_sections]}
    finally:
        db.close()

    print(json.dumps(evidence, indent=2))


def upload_file(path: Path) -> dict:
    with path.open("rb") as file_obj:
        response = requests.post(f"{API}/uploads/rfp", files={"file": (path.name, file_obj)}, timeout=120)
    try:
        return response.json()
    except Exception:
        return {"status_code": response.status_code, "text": response.text}


def chat(rfp_id: int, question: str) -> dict:
    return requests.post(f"{API}/private-rfp/{rfp_id}/chat", json={"question": question}, timeout=240).json()


def summarize_chat(result: dict) -> dict:
    return {
        "provider": result.get("provider"),
        "model_used": result.get("model_used"),
        "retrieval_mode": result.get("retrieval_mode"),
        "external_api_used": result.get("external_api_used"),
        "source_chunk_count": len(result.get("source_chunks", [])),
        "answer_preview": str(result.get("answer", ""))[:240],
    }


def duplicate_heading_count(sections: list[dict]) -> int:
    count = 0
    for section in sections:
        title = section["section_title"].strip().lower()
        lines = [line.strip().lower() for line in section["section_content"].splitlines()]
        count += sum(1 for line in lines if line == title)
    return count


def docx_contains(path: Path, phrase: str) -> bool:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return phrase in text


def create_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Searchable RFP for Private Validation", level=1)
    for paragraph in valid_rfp_text().split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(path)


def valid_rfp_text() -> str:
    base = """
Request for Proposal: Local Private RFP Analysis Platform

The buyer seeks a secure local-first document analysis platform to support upload, parsing, classification, retrieval, private chat, short draft generation, and DOCX export. The scope includes PDF, DOCX, and TXT ingestion, deterministic document quality checks, local chunking, local embeddings using Ollama, and private question answering grounded in retrieved excerpts.

The selected vendor must provide a professional application that runs on a Windows laptop without dedicated GPU resources. The system must keep all RFP content on the machine and must not transmit documents, chunks, embeddings, or prompts to external cloud AI providers. The platform should expose a clean backend API and a minimal frontend for dashboard, upload, chat, and settings.

Functional requirements include local upload validation, text extraction, chunk creation, lexical retrieval fallback, semantic retrieval when local embeddings are available, private chat responses, source chunk references, draft generation for eight short sections, and DOCX export with a privacy note. Technical requirements include SQLite storage for the MVP, FastAPI backend services, Ollama llama3.2:3b for local chat, and nomic-embed-text for local embeddings.

Security requirements include private mode operation, no external AI providers, local-only processing, no arbitrary shell execution, no arbitrary file access, and clear human review requirements for generated drafts. The tool must support auditability and controlled behavior suitable for enterprise evaluation.

Implementation should be completed in phases. The first phase covers backend foundation, upload, parsing, classification, chunking, retrieval, and private chat. The second phase covers local embeddings, hybrid retrieval, private short draft generation, DOCX export, and a local MCP wrapper. The expected deadline for MVP validation is 30 June 2026, with a pilot review planned after local testing.

Submission requirements include a working local backend, clear API documentation, test scripts, and evidence that external_api_used remains false. The proposal should describe the approach, constraints, assumptions, risks, mitigations, implementation plan, and support for future controlled integrations.
"""
    return "\n\n".join([base.strip()] * 3)


if __name__ == "__main__":
    main()
