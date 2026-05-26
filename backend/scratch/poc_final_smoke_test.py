from __future__ import annotations

import subprocess
import sys
from pathlib import Path
from tempfile import TemporaryDirectory

import fitz
from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))

from app.core.database import SessionLocal, init_db
from app.models.rfp import AuditLog, LocalUsageLog, RFPChunk, RFPDocument, RFPDraftSection
from app.services.classification_service import classify_document
from app.services.docx_export_service import export_private_short_draft
from app.services.local_embedding_service import embed_chunks_for_rfp
from app.services.local_llm_service import check_ollama_health
from app.services.metadata_extraction_service import extract_probable_metadata
from app.services.parsing_service import extract_text_from_file
from app.services.private_chat_service import answer_private_rfp_question
from app.services.private_generation_service import generate_private_short_draft
from app.services.retrieval_service import search_chunks
from app.services.ollama_health_service import get_ollama_health
from app.services.usage_service import usage_summary
from mcp_server.tools.generation_tools import generate_private_short_draft_tool
from mcp_server.tools.rfp_tools import list_rfps_tool

rows: list[tuple[str, str, str, str]] = []


def add(test: str, ok: bool, evidence: str, notes: str = "") -> None:
    rows.append((test, "PASS" if ok else "FAIL", evidence, notes))


def create_tiny_doc(path: Path) -> Path:
    path.write_text("Short RFP note.", encoding="utf-8")
    return path


def create_large_pdf(path: Path) -> Path:
    doc = fitz.open()
    page_text = (
        "Request for Proposal: Secure Local RFP Intelligence Workspace\n\n"
        "Client: Enterprise Procurement Department\n"
        "Submission deadline: 30 June 2026\n"
        "The proposal should remain local, private, and auditable.\n"
        "The buyer expects local chat, retrieval, draft generation, and DOCX export.\n"
        "This page contains enough text to support private RFP analysis.\n"
    )
    for page_number in range(25):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {page_number + 1}\n\n{page_text * 2}")
    doc.save(path)
    doc.close()
    return path


def create_uploaded_rfp(file_path: Path) -> int:
    parsed = extract_text_from_file(str(file_path))
    classification = classify_document(
        int(parsed["page_count"]),
        int(parsed["word_count"]),
        int(parsed["character_count"]),
        str(parsed["text"]),
    )
    metadata = extract_probable_metadata(str(parsed["text"]), file_path.name)
    db = SessionLocal()
    try:
        rfp = RFPDocument(
            title=metadata["probable_title"] or file_path.stem,
            original_filename=file_path.name,
            stored_filename=file_path.name,
            file_path=str(file_path),
            file_type=file_path.suffix.lstrip("."),
            file_size=file_path.stat().st_size,
            page_count=int(parsed["page_count"]),
            word_count=int(parsed["word_count"]),
            character_count=int(parsed["character_count"]),
            line_count=int(parsed["line_count"]),
            extracted_text=str(parsed["text"]),
            probable_title=metadata["probable_title"],
            probable_client=metadata["probable_client"],
            probable_deadline=metadata["probable_deadline"],
            probable_submission_date=metadata["probable_submission_date"],
            document_quality=classification["document_quality"],
            classification_reason=classification["reason"],
            status=classification["status"],
        )
        db.add(rfp)
        db.commit()
        db.refresh(rfp)
        if rfp.document_quality in {"valid_rfp", "limited_but_valid"}:
            from app.services.chunking_service import create_chunks_for_rfp

            create_chunks_for_rfp(db, rfp.id, rfp.extracted_text or "")
        return rfp.id
    finally:
        db.close()


def main() -> None:
    init_db()
    ollama = check_ollama_health()
    health = get_ollama_health()
    add("Backend health", bool(ollama.get("ok")), "FastAPI startup and DB init complete", "")
    add("Ollama available", bool(health.get("available")), str(health), "")
    add("Chat model available", bool(health.get("chat_model_available")), str(health.get("models")), "")
    add("Embedding model available", bool(health.get("embedding_model_available")), str(health.get("models")), "")

    with TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        tiny_path = create_tiny_doc(tmp_path / "tiny_rfp.txt")
        tiny_rfp_id = create_uploaded_rfp(tiny_path)
        db = SessionLocal()
        try:
            tiny = db.get(RFPDocument, tiny_rfp_id)
            add("Tiny upload classified", tiny is not None and tiny.status == "needs_more_detail", f"status={tiny.status if tiny else None}", "")
        finally:
            db.close()

        demo_pdf = create_large_pdf(tmp_path / "demo_25_page_rfp.pdf")
        rfp_id = create_uploaded_rfp(demo_pdf)

        db = SessionLocal()
        try:
            rfp = db.get(RFPDocument, rfp_id)
            add(
                "Large upload classified",
                rfp is not None and rfp.status in {"ready_for_private_chat", "extraction_needs_review"},
                f"rfp_id={rfp_id}, status={rfp.status if rfp else None}, pages={rfp.page_count if rfp else None}, words={rfp.word_count if rfp else None}",
                "",
            )
            add("Extracted preview exists", bool(rfp and (rfp.extracted_text or "")[:500]), (rfp.extracted_text or "")[:120] if rfp else "", "")
            chunks = db.query(RFPChunk).filter(RFPChunk.rfp_id == rfp_id).count()
            add("Chunk count > 0", chunks > 0, f"chunks={chunks}", "")
            emb = embed_chunks_for_rfp(db, rfp_id)
            add("Embedding generation", emb["embedded_chunks"] > 0 and emb["failed_chunks"] == 0, str(emb), "")
            retrieval = search_chunks(db, rfp_id, "scope requirements deadline", mode="hybrid")
            add("Retrieval returns chunks", len(retrieval) > 0, f"count={len(retrieval)}, type={retrieval[0].get('retrieval_type') if retrieval else None}", "")
            hi = answer_private_rfp_question(db, rfp_id, "hi")
            add("Greeting local", hi.get("provider") == "local", str(hi.get("provider")), "")
            chat = answer_private_rfp_question(db, rfp_id, "What is this RFP about?")
            add("Chat local_ollama", chat.get("provider") == "local_ollama", str(chat.get("provider")), "")
            add("external_api_used false", chat.get("external_api_used") is False, str(chat.get("external_api_used")), "")
            partial_guard = generate_private_short_draft_tool(rfp_id, confirm=False)
            add("MCP confirm guard", "confirmation" in str(partial_guard).lower(), str(partial_guard), "")
            gen = generate_private_short_draft(db, rfp_id, confirm_regenerate=True)
            sections = db.query(RFPDraftSection).filter(RFPDraftSection.rfp_id == rfp_id).count()
            add("Draft response", gen.get("status") in {"completed", "failed_partial"}, str(gen), "")
            add("Short draft section count", sections <= 8, f"sections={sections}, response={gen}", "")
            valid_sections = db.query(RFPDraftSection).filter(RFPDraftSection.rfp_id == rfp_id).all()
            invalid_text = any(
                any(marker in (section.section_content or "").lower() for marker in ["ollama is not running", "please start ollama", "traceback", "exception"])
                for section in valid_sections
            )
            add("No infra text saved", not invalid_text, f"section_count={len(valid_sections)}", "")
            job = db.query(AuditLog).filter(AuditLog.rfp_id == rfp_id).count()
            add("Audit logs exist", job > 0, f"audit_count={job}", "")
            usage = db.query(LocalUsageLog).filter(LocalUsageLog.rfp_id == rfp_id).count()
            add("Usage logs exist", usage > 0, f"usage_count={usage}", "")
            export = export_private_short_draft(db, rfp_id)
            path = Path(export["file_path"])
            add("DOCX export", path.exists() and path.stat().st_size > 0, f"{path}, size={path.stat().st_size if path.exists() else 0}", "")
            doc = Document(path)
            doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            blocked = any(marker in doc_text.lower() for marker in ["ollama is not running", "please start ollama", "traceback", "exception"])
            add("DOCX content clean", not blocked, path.name, "")
            add("Usage summary", usage_summary(db)["total_calls"] > 0, str(usage_summary(db)), "")
            add("MCP list works", len(list_rfps_tool()) > 0, f"rfps={len(list_rfps_tool())}", "")
        finally:
            db.close()

    scan_script = Path(__file__).with_name("security_scan_private_mode.py")
    scan = subprocess.run([sys.executable, str(scan_script)], capture_output=True, text=True)
    add("Security scan", scan.returncode == 0, scan.stdout.strip().splitlines()[-1] if scan.stdout else "", scan.stderr.strip())

    print("Test | Status | Evidence | Notes")
    print("---|---|---|---")
    for row in rows:
        print(" | ".join(str(col).replace("\n", " ") for col in row))

    if any(row[1] == "FAIL" for row in rows):
        sys.exit(1)


if __name__ == "__main__":
    main()
